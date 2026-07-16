"""DOCX/PDF/text parsers with explicit OCR and table-preservation boundaries."""

from __future__ import annotations

import json
from pathlib import Path
import re
import unicodedata
from typing import Any, Protocol, Sequence

from ingest.models import DocumentElement, ElementKind, ParsedDocument


class UnsupportedDocumentError(ValueError):
    """Raised when an input must be converted or unpacked before ingestion."""


class OCRRequiredError(ValueError):
    """Raised instead of silently accepting an image-only PDF."""


class OCRProvider(Protocol):
    def pages(self, pdf_path: Path, *, expected_pages: int) -> list[str]: ...


class SidecarOCRProvider:
    """Read reviewed OCR pages from ``<filename>.ocr.json`` sidecars."""

    def __init__(self, directory: str | Path) -> None:
        self.directory = Path(directory)

    def sidecar_path(self, pdf_path: Path) -> Path:
        direct = self.directory / f"{pdf_path.name}.ocr.json"
        if direct.is_file():
            return direct
        return self.directory / f"{pdf_path.stem}.ocr.json"

    def pages(self, pdf_path: Path, *, expected_pages: int) -> list[str]:
        sidecar = self.sidecar_path(pdf_path)
        if not sidecar.is_file():
            raise OCRRequiredError(
                f"OCR sidecar is missing for {pdf_path.name}: expected {sidecar}"
            )
        try:
            payload = json.loads(sidecar.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError) as exc:
            raise ValueError(f"OCR sidecar is unreadable: {sidecar}") from exc
        page_map = self.page_map(pdf_path, expected_pages=expected_pages, payload=payload)
        if set(page_map) != set(range(1, expected_pages + 1)):
            raise ValueError(
                f"OCR page count mismatch for {pdf_path.name}: "
                f"found {len(page_map)}, expected {expected_pages}"
            )
        return [page_map[page] for page in range(1, expected_pages + 1)]

    def page_map(
        self,
        pdf_path: Path,
        *,
        expected_pages: int,
        payload: dict[str, Any] | None = None,
    ) -> dict[int, str]:
        sidecar = self.sidecar_path(pdf_path)
        if payload is None:
            if not sidecar.is_file():
                raise OCRRequiredError(
                    f"OCR sidecar is missing for {pdf_path.name}: expected {sidecar}"
                )
            try:
                payload = json.loads(sidecar.read_text(encoding="utf-8"))
            except (OSError, json.JSONDecodeError) as exc:
                raise ValueError(f"OCR sidecar is unreadable: {sidecar}") from exc
        pages = payload.get("pages") if isinstance(payload, dict) else None
        if not isinstance(pages, list) or not pages:
            raise ValueError(f"OCR sidecar must contain a non-empty pages list: {sidecar}")
        result: dict[int, str] = {}
        for item in pages:
            page = item.get("page") if isinstance(item, dict) else None
            if not isinstance(page, int) or not 1 <= page <= expected_pages or page in result:
                raise ValueError(
                    f"OCR pages must be unique and within the PDF page range: {sidecar}"
                )
            text = item.get("text")
            if not isinstance(text, str) or not normalize_text(text):
                raise ValueError(
                    f"OCR page {page} is empty or invalid: {sidecar}"
                )
            result[page] = text
        return result


_CJK_SPACE_RE = re.compile(r"(?<=[\u3400-\u9fff])[ \t\u3000]+(?=[\u3400-\u9fff])")
_INLINE_SPACE_RE = re.compile(r"[ \t\u3000]+")
_BLANK_LINES_RE = re.compile(r"\n{3,}")
_CJK_RADICAL_TRANSLATION = str.maketrans({"⻚": "页"})
_HEADING_RE = re.compile(
    r"^(?:第[一二三四五六七八九十百千万零〇两0-9]+[章节条]|"
    r"[一二三四五六七八九十百]+、|附(?:表|件)\s*[一二三四五六七八九十0-9]+)"
)
_LIST_START_RE = re.compile(r"^(?:\d{1,3}\s*[.·、]|[（(][一二三四五六七八九十0-9]+[）)])")
_PAGE_MARK_RE = re.compile(r"^[—\-–]?\s*\d+(?:\s*/\s*\d+)?\s*[—\-–]?$")
_WEB_PRINT_RE = re.compile(r"^\d{4}/\d{1,2}/\d{1,2}\s+\d{1,2}:\d{2}\b")
_WEB_PRINT_SITE_SUFFIX = "-\u897f\u5357\u8d22\u7ecf\u5927\u5b66\u8ba1\u7b97\u673a\u4e0e\u4eba\u5de5\u667a\u80fd\u5b66\u9662"
_WEB_PRINT_FOOTER = "\u7248\u6743\u6240\u6709@ \u897f\u5357\u8d22\u7ecf\u5927\u5b66"
_WEB_PRINT_URL_RE = re.compile(
    r"https://it\.swufe\.edu\.cn/info/\d+/\d+\.htm\s+\d+/\d+\s*$")


def normalize_text(value: str) -> str:
    text = unicodedata.normalize("NFKC", value)
    text = text.translate(_CJK_RADICAL_TRANSLATION)
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    text = _CJK_SPACE_RE.sub("", text)
    lines = [_INLINE_SPACE_RE.sub(" ", line).strip() for line in text.splitlines()]
    return _BLANK_LINES_RE.sub("\n\n", "\n".join(lines)).strip()


def join_wrapped_lines(value: str) -> str:
    """Join visual line wraps while keeping headings and numbered items separate."""

    lines = [line for line in normalize_text(value).splitlines() if line]
    logical: list[str] = []
    for line in lines:
        if _PAGE_MARK_RE.fullmatch(line):
            continue
        starts_structure = _looks_like_heading(line) or bool(_LIST_START_RE.match(line))
        previous_is_heading = bool(logical and _is_standalone_heading(logical[-1]))
        if not logical or starts_structure or previous_is_heading:
            logical.append(line)
            continue
        separator = " " if logical[-1][-1:].isascii() and line[:1].isascii() else ""
        logical[-1] += separator + line
    return "\n".join(logical)


def _looks_like_heading(text: str, style_name: str = "") -> bool:
    style = style_name.lower()
    if style.startswith("heading") or style in {"title", "subtitle", "标题", "副标题"}:
        return True
    return len(text) <= 90 and bool(_HEADING_RE.match(text))


def _is_standalone_heading(text: str) -> bool:
    match = _HEADING_RE.match(text)
    return bool(match and len(text[match.end() :].strip()) <= 20)


def table_to_markdown(rows: list[list[Any]]) -> str:
    normalized: list[list[str]] = []
    width = 0
    for row in rows:
        cells = [
            normalize_text("" if cell is None else str(cell)).replace("|", "\\|")
            for cell in row
        ]
        width = max(width, len(cells))
        normalized.append(cells)
    normalized = [row + [""] * (width - len(row)) for row in normalized if any(row)]
    if not normalized or width == 0:
        return ""
    header = normalized[0]
    if not any(header):
        header = [f"列{index}" for index in range(1, width + 1)]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
    return "\n".join(lines)


def _parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX parsing; install requirements-ingest.txt"
        ) from exc

    document = Document(path)
    elements: list[DocumentElement] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            text = normalize_text(paragraph.text)
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            kind: ElementKind = (
                "heading" if _looks_like_heading(text, style_name) else "paragraph"
            )
            elements.append(DocumentElement(kind, text))
        elif isinstance(child, CT_Tbl):
            table = Table(child, document)
            markdown = table_to_markdown(
                [[cell.text for cell in row.cells] for row in table.rows]
            )
            if markdown:
                elements.append(DocumentElement("table", markdown))
    if not elements:
        raise ValueError(f"document contains no extractable text or tables: {path}")
    warnings: list[str] = []
    if document.inline_shapes:
        warnings.append(
            f"contains {len(document.inline_shapes)} inline images; image text is not OCRed"
        )
    return ParsedDocument(path, elements, warnings=warnings)


def _needs_ocr(page_texts: Sequence[str], *, minimum_chars_per_page: int = 80) -> bool:
    meaningful = sum(
        len(re.findall(r"[\u3400-\u9fffA-Za-z0-9]", text)) for text in page_texts
    )
    return meaningful < max(160, len(page_texts) * minimum_chars_per_page)


def _is_web_print_page(text: str) -> bool:
    return bool(_WEB_PRINT_RE.match(text))


def _clean_web_print_paragraph(text: str) -> str:
    cleaned = _WEB_PRINT_RE.sub("", text, count=1).lstrip()
    if _WEB_PRINT_SITE_SUFFIX in cleaned:
        cleaned = cleaned.split(_WEB_PRINT_SITE_SUFFIX, 1)[1]
    if _WEB_PRINT_FOOTER in cleaned:
        cleaned = cleaned.split(_WEB_PRINT_FOOTER, 1)[0]
    cleaned = _WEB_PRINT_URL_RE.sub("", cleaned)
    return normalize_text(cleaned)


def _clean_web_print_table(markdown: str) -> str:
    return markdown.replace("\u5206\u4eab", "")


def _parse_pdf(path: Path, ocr_provider: OCRProvider | None) -> ParsedDocument:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError(
            "pdfplumber is required for PDF parsing; install requirements-ingest.txt"
        ) from exc

    page_texts: list[str] = []
    page_tables: list[list[list[list[object]]]] = []
    page_has_images: list[bool] = []
    with pdfplumber.open(path) as document:
        page_count = len(document.pages)
        for page in document.pages:
            page_texts.append(join_wrapped_lines(page.extract_text() or ""))
            page_has_images.append(bool(page.images))
            try:
                page_tables.append(page.extract_tables() or [])
            except Exception:
                page_tables.append([])

    if _needs_ocr(page_texts):
        if ocr_provider is None:
            raise OCRRequiredError(
                f"PDF has insufficient embedded text and requires OCR: {path.name}"
            )
        pages = ocr_provider.pages(path, expected_pages=page_count)
        elements = [
            DocumentElement("paragraph", join_wrapped_lines(text), page=index)
            for index, text in enumerate(pages, start=1)
        ]
        return ParsedDocument(path, elements, page_count, ["ocr_used"])

    image_only_pages = [
        page_number
        for page_number, (text, tables, has_images) in enumerate(
            zip(page_texts, page_tables, page_has_images), start=1
        )
        if not text and not tables and has_images
    ]
    partial_ocr: dict[int, str] = {}
    if image_only_pages:
        if ocr_provider is None or not hasattr(ocr_provider, "page_map"):
            raise OCRRequiredError(
                f"PDF contains image-only pages that require OCR: {path.name} "
                f"pages {image_only_pages}"
            )
        partial_ocr = ocr_provider.page_map(path, expected_pages=page_count)  # type: ignore[attr-defined]
        missing = sorted(set(image_only_pages) - set(partial_ocr))
        if missing:
            raise OCRRequiredError(
                f"OCR sidecar is missing image-only pages for {path.name}: {missing}"
            )

    elements: list[DocumentElement] = []
    web_print_cleaned = False
    for page_number, (text, tables) in enumerate(zip(page_texts, page_tables), start=1):
        web_print_page = _is_web_print_page(text)
        if text and not web_print_page:
            elements.append(DocumentElement("paragraph", text, page=page_number))
        elif page_number in partial_ocr:
            elements.append(
                DocumentElement("paragraph", join_wrapped_lines(partial_ocr[page_number]), page=page_number)
            )
        if web_print_page:
            web_print_cleaned = True
            if not tables:
                cleaned = _clean_web_print_paragraph(text)
                if cleaned:
                    elements.append(DocumentElement("paragraph", cleaned, page=page_number))
        for rows in tables:
            markdown = table_to_markdown(rows)
            if web_print_page:
                markdown = _clean_web_print_table(markdown)
            if markdown:
                elements.append(DocumentElement("table", markdown, page=page_number))
    if not elements:
        raise ValueError(f"PDF contains no extractable content: {path}")
    warnings = ["web_print_noise_removed"] if web_print_cleaned else []
    if partial_ocr:
        warnings.append("partial_ocr_used")
    return ParsedDocument(path, elements, page_count, warnings)


def _parse_text(path: Path) -> ParsedDocument:
    text = normalize_text(path.read_text(encoding="utf-8-sig"))
    if not text:
        raise ValueError(f"text source is empty: {path}")
    return ParsedDocument(path, [DocumentElement("paragraph", text)])


def parse_document(
    path: str | Path,
    *,
    ocr_provider: OCRProvider | None = None,
) -> ParsedDocument:
    source = Path(path)
    if not source.is_file():
        raise FileNotFoundError(source)
    suffix = source.suffix.lower()
    if suffix == ".docx":
        return _parse_docx(source)
    if suffix == ".pdf":
        return _parse_pdf(source, ocr_provider)
    if suffix in {".txt", ".md"}:
        return _parse_text(source)
    if suffix == ".doc":
        raise UnsupportedDocumentError(
            f"legacy DOC must be converted to DOCX before ingestion: {source.name}"
        )
    if suffix == ".zip":
        raise UnsupportedDocumentError(
            f"ZIP archives must be unpacked and registered per document: {source.name}"
        )
    raise UnsupportedDocumentError(f"unsupported document type: {suffix or '<none>'}")
