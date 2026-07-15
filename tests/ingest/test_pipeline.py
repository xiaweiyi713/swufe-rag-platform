from __future__ import annotations

import csv
from pathlib import Path
import tempfile
import unittest

from docx import Document

from ingest.pipeline import ingest_sources
from ingest.sources import SOURCE_FIELDS
from retrieval.index import load_chunks


class IngestionPipelineTests(unittest.TestCase):
    def test_docx_to_frozen_chunks_delivery(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            raw = root / "raw" / "school"
            raw.mkdir(parents=True)
            document_path = raw / "policy.docx"
            document = Document()
            document.add_heading("第一章 总则", level=1)
            document.add_paragraph("第一条 本规定适用于本科生。")
            document.save(document_path)

            sources = root / "sources.csv"
            with sources.open("w", encoding="utf-8", newline="") as handle:
                writer = csv.DictWriter(handle, fieldnames=SOURCE_FIELDS)
                writer.writeheader()
                writer.writerow(
                    {
                        "file": "school/policy.docx",
                        "doc_title": "本科测试规定",
                        "level": "校级",
                        "college": "全校",
                        "cohort": "不限",
                        "year": "2026",
                        "status": "现行",
                        "page_url": "https://jwc.swufe.edu.cn/policy",
                        "file_url": "https://jwc.swufe.edu.cn/policy.docx",
                        "collected_at": "2026-07-14",
                    }
                )

            output = root / "chunks.jsonl"
            report = ingest_sources(
                sources,
                root / "raw",
                output,
                report_path=root / "report.json",
            )
            chunks = load_chunks(output)
            self.assertEqual(report["source_count"], 1)
            self.assertEqual(report["chunk_count"], len(chunks))
            self.assertEqual(chunks[0]["doc_title"], "本科测试规定")


if __name__ == "__main__":
    unittest.main()
