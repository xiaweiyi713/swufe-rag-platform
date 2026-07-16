from __future__ import annotations

import json
from pathlib import Path
import tempfile
import unittest

from ingest.parse import (
    _clean_web_print_paragraph,
    _clean_web_print_table,
    _is_web_print_page,
    SidecarOCRProvider,
    join_wrapped_lines,
    normalize_text,
    parse_document,
    table_to_markdown,
)


class ParseTests(unittest.TestCase):
    def test_cjk_compatibility_forms_are_normalized(self) -> None:
        self.assertEqual(normalize_text("西南财经⼤学主⻚ 2026年06⽉18⽇"), "西南财经大学主页 2026年06月18日")

    def test_windows_ocr_cjk_spacing_is_normalized(self) -> None:
        self.assertEqual(normalize_text("西 南 财 经 大 学\n第 一 条"), "西南财经大学\n第一条")

    def test_table_is_rendered_as_complete_markdown(self) -> None:
        markdown = table_to_markdown([["课程", "学分"], ["人工智能", "3"]])
        self.assertIn("| 课程 | 学分 |", markdown)
        self.assertIn("| 人工智能 | 3 |", markdown)

    def test_web_print_noise_is_detected_and_removed(self) -> None:
        self.assertTrue(
            _is_web_print_page("2026/7/15 12:34 SWUFE")
        )
        self.assertFalse(_is_web_print_page("SWUFE"))
        paragraph = (
            "2026/7/15 12:34 title"
            "-\u897f\u5357\u8d22\u7ecf\u5927\u5b66\u8ba1\u7b97\u673a\u4e0e\u4eba\u5de5\u667a\u80fd\u5b66\u9662"
            "\nbody text https://it.swufe.edu.cn/info/1166/1.htm 2/3"
        )
        self.assertEqual(_clean_web_print_paragraph(paragraph), "body text")
        noisy = "\u653b\u5206\u4eab\u8bfb\u897f\u5357\u8d22\u7ecf\u5206\u4eab\u5927\u5b66"
        self.assertEqual(_clean_web_print_table(noisy), "\u653b\u8bfb\u897f\u5357\u8d22\u7ecf\u5927\u5b66")

    def test_visual_line_wraps_join_without_losing_numbered_items(self) -> None:
        text = "第一条 毕业最低学分要求：计算机科学与技术151学分、人工智能150学分、网络\n空间安全150学分。\n1 · 第一项。\n2 · 第二项。\n— 1 —"
        joined = join_wrapped_lines(text)
        self.assertIn("网络空间安全150学分", joined)
        self.assertIn("\n1 · 第一项。\n2 · 第二项。", joined)
        self.assertNotIn("— 1 —", joined)

    def test_docx_preserves_body_order_and_table(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "policy.docx"
            document = Document()
            document.add_heading("第一章 总则", level=1)
            document.add_paragraph("第一条 本规定适用于本科生。")
            table = document.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "课程"
            table.rows[0].cells[1].text = "学分"
            table.rows[1].cells[0].text = "人工智能"
            table.rows[1].cells[1].text = "3"
            document.save(path)

            parsed = parse_document(path)
            self.assertEqual([item.kind for item in parsed.elements], ["heading", "heading", "table"])
            self.assertIn("人工智能", parsed.elements[-1].text)

    def test_sidecar_requires_consecutive_pages_and_exact_count(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            pdf = root / "scan.pdf"
            pdf.write_bytes(b"unused")
            sidecar = root / "scan.pdf.ocr.json"
            sidecar.write_text(
                json.dumps({"pages": [{"page": 1, "text": "第一页"}]}, ensure_ascii=False),
                encoding="utf-8",
            )
            provider = SidecarOCRProvider(root)
            self.assertEqual(provider.pages(pdf, expected_pages=1), ["第一页"])
            with self.assertRaisesRegex(ValueError, "page count mismatch"):
                provider.pages(pdf, expected_pages=2)


    def test_sidecar_page_map_accepts_reviewed_partial_pages(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            pdf = root / "mixed.pdf"
            pdf.write_bytes(b"unused")
            sidecar = root / "mixed.pdf.ocr.json"
            sidecar.write_text(
                json.dumps(
                    {"pages": [{"page": 1, "text": "图片封面标题"}]},
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            provider = SidecarOCRProvider(root)
            self.assertEqual(provider.page_map(pdf, expected_pages=10), {1: "图片封面标题"})

if __name__ == "__main__":
    unittest.main()
